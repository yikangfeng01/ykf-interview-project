"""Flask web application for signature page generation system."""

import os
import io
import sys
import datetime as dt

from flask import Flask, request, jsonify, send_from_directory, send_file
from flask.json.provider import DefaultJSONProvider
from typing import Optional

from auth.service import authenticate, get_current_user, logout as auth_logout
from auth.db import init_database, get_connection
from data import projects, templates, signature_variables, project_templates, project_variables, public_variables

app = Flask(__name__, static_folder="static", static_url_path="")


class ISOJSONProvider(DefaultJSONProvider):
    """Custom JSON provider that serializes datetime as ISO format instead of HTTP-date."""
    @staticmethod
    def default(o):
        if isinstance(o, (dt.datetime, dt.date)):
            return o.isoformat()
        return DefaultJSONProvider.default(o)


app.json = ISOJSONProvider(app)

# ---- Database initialization ----
if os.environ.get("SKIP_DB_INIT") not in ("1", "true", "yes"):
    try:
        init_database()
    except RuntimeError as e:
        print(f"[app] FATAL: {e}", file=sys.stderr)
        sys.exit(1)
else:
    print("[app] SKIP_DB_INIT is set; skipping database initialization")


def _extract_token() -> Optional[str]:
    """Extract session token from request.
    
    Checks in order: JSON body > form data > query param > Authorization header.
    """
    # 1. JSON body
    data = request.get_json(silent=True)
    if data and data.get("token"):
        return data["token"].strip()

    # 2. Form data (for multipart/form-data requests)
    token = request.form.get("token")
    if token:
        return token.strip()

    # 3. Query parameter (for GET requests)
    token = request.args.get("token")
    if token:
        return token.strip()

    # 4. Authorization header
    auth = request.headers.get("Authorization", "")
    if auth.startswith("Bearer "):
        return auth[7:].strip()

    return None


def _require_auth():
    """Verify authentication. Returns (user, error_response) tuple.
    
    If authenticated, returns (user, None).
    If not, returns (None, error_json_response).
    """
    token = _extract_token()
    if not token:
        return None, (jsonify({"error": "Token is required"}), 401)

    user = get_current_user(token)
    if user is None:
        return None, (jsonify({"error": "Invalid token"}), 401)

    return user, None


# ============ Static ============

@app.route("/")
def index():
    """Serve the SPA shell."""
    return send_from_directory("static", "index.html")


# ============ Auth API ============

@app.route("/api/login", methods=["POST"])
def api_login():
    """Login API endpoint.
    
    Expects JSON: {"username": "...", "password": "..."}
    Returns JSON: {"token": "..."} on success, {"error": "..."} on failure.
    """
    data = request.get_json(silent=True)
    if data is None:
        return jsonify({"error": "Invalid JSON body"}), 400

    username = (data.get("username") or "").strip()
    password = data.get("password") or ""

    success, result = authenticate(username, password)

    if success:
        return jsonify({"token": result}), 200
    else:
        if "required" in result.lower():
            return jsonify({"error": result}), 400
        else:
            return jsonify({"error": result}), 401


@app.route("/api/verify", methods=["POST"])
def api_verify():
    """Verify a session token.

    Expects JSON: {"token": "..."}
    Returns JSON: {"username": "..."} on success, {"error": "..."} on failure.
    """
    data = request.get_json(silent=True)
    if data is None:
        return jsonify({"error": "Invalid JSON body"}), 400

    token = (data.get("token") or "").strip()
    if not token:
        return jsonify({"error": "Token is required"}), 400

    user = get_current_user(token)
    if user is None:
        return jsonify({"error": "Invalid token"}), 401

    return jsonify({"username": user.username}), 200


@app.route("/api/logout", methods=["POST"])
def api_logout():
    """Logout API endpoint.

    Expects JSON: {"token": "..."}
    Returns JSON: {"message": "ok"} on success, {"error": "..."} on failure.
    """
    data = request.get_json(silent=True)
    if data is None:
        return jsonify({"error": "Invalid JSON body"}), 400

    token = (data.get("token") or "").strip()
    if not token:
        return jsonify({"error": "Token is required"}), 400

    try:
        auth_logout(token)
        return jsonify({"message": "ok"}), 200
    except ValueError as e:
        return jsonify({"error": str(e)}), 400


# ============ Projects API ============

@app.route("/api/projects", methods=["GET", "POST"])
def api_projects():
    """List projects (GET) or create a new project (POST)."""
    user, err = _require_auth()
    if err:
        return err

    if request.method == "GET":
        result = projects.list_by_user(user.username)
        return jsonify(result), 200

    # POST: create project
    data = request.get_json(silent=True)
    if data is None:
        return jsonify({"error": "Invalid JSON body"}), 400

    name = (data.get("name") or "").strip()
    if not name:
        return jsonify({"error": "Project name is required"}), 400

    description = (data.get("description") or "").strip() or None
    project = projects.create(name, user.username, description=description)
    return jsonify(project), 201


@app.route("/api/projects/<int:project_id>", methods=["GET", "PUT", "DELETE"])
def api_project(project_id):
    """Get, update, or delete a specific project."""
    user, err = _require_auth()
    if err:
        return err

    project = projects.get_by_id(project_id)
    if project is None:
        return jsonify({"error": "Project not found"}), 404

    # Ownership check
    if project["created_by"] != user.username:
        return jsonify({"error": "Forbidden"}), 403

    if request.method == "GET":
        return jsonify(project), 200

    if request.method == "PUT":
        data = request.get_json(silent=True)
        if data is None:
            return jsonify({"error": "Invalid JSON body"}), 400

        updates = {}
        name = (data.get("name") or "").strip()
        if name:
            updates["name"] = name
        if "description" in data:
            desc = (data.get("description") or "").strip()
            updates["description"] = desc if desc else None
        if "status" in data:
            updates["status"] = data["status"]

        updated = projects.update(project_id, **updates)
        return jsonify(updated), 200

    # DELETE
    projects.delete(project_id)
    return "", 204


# ============ Templates API ============

UPLOAD_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "uploads", "templates")


def _ensure_upload_dir():
    """Ensure the upload directory exists."""
    os.makedirs(UPLOAD_DIR, exist_ok=True)


@app.route("/api/templates", methods=["GET", "POST"])
def api_templates():
    """List templates (GET) or upload a new template (POST)."""
    user, err = _require_auth()
    if err:
        return err

    if request.method == "GET":
        category = request.args.get("category")
        result = templates.list_all(category=category)
        return jsonify(result), 200

    # POST: upload template with DOCX file
    if "file" not in request.files:
        return jsonify({"error": "DOCX file is required"}), 400

    docx_file = request.files["file"]
    if docx_file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    if not docx_file.filename.lower().endswith(".docx"):
        return jsonify({"error": "Only DOCX files are allowed"}), 400

    name = (request.form.get("name") or "").strip()
    if not name:
        return jsonify({"error": "Template name is required"}), 400

    category = (request.form.get("category") or "").strip()
    description = (request.form.get("description") or "").strip()

    # Insert record first to get the ID
    tmpl = templates.create(name=name, category=category, description=description)

    # Save DOCX file
    _ensure_upload_dir()
    tmpl_dir = os.path.join(UPLOAD_DIR, str(tmpl["id"]))
    os.makedirs(tmpl_dir, exist_ok=True)
    safe_filename = "template.docx"
    relative_path = f"uploads/templates/{tmpl['id']}/{safe_filename}"
    file_path = os.path.join(tmpl_dir, safe_filename)
    docx_file.save(file_path)

    # Update file_path in DB
    updated = templates.update(tmpl["id"], file_path=relative_path)
    return jsonify(updated), 201


@app.route("/api/templates/<int:template_id>", methods=["PUT", "DELETE"])
def api_template(template_id):
    """Update or delete a specific template."""
    user, err = _require_auth()
    if err:
        return err

    tmpl = templates.get_by_id(template_id)
    if tmpl is None:
        return jsonify({"error": "Template not found"}), 404

    if request.method == "PUT":
        data = request.get_json(silent=True)
        if data is None:
            return jsonify({"error": "Invalid JSON body"}), 400

        name = (data.get("name") or "").strip() or None
        category = (data.get("category") or "").strip() or None
        description = (data.get("description") or "").strip() or None
        file_path = (data.get("file_path") or "").strip() or None

        updated = templates.update(template_id, name=name, category=category,
                                   description=description, file_path=file_path)
        if updated is None:
            return jsonify({"error": "Template not found"}), 404

        # Cascade update to project_templates bound to this public template
        cascade_set = []
        cascade_vals = []
        if name is not None:
            cascade_set.append("name = %s")
            cascade_vals.append(name)
        if category is not None:
            cascade_set.append("category = %s")
            cascade_vals.append(category)
        if description is not None:
            cascade_set.append("description = %s")
            cascade_vals.append(description)
        if file_path is not None:
            cascade_set.append("file_path = %s")
            cascade_vals.append(file_path)
        if cascade_set:
            cascade_vals.append(template_id)
            conn = get_connection()
            cur = conn.cursor()
            try:
                cur.execute(
                    f"UPDATE project_templates SET {', '.join(cascade_set)} "
                    "WHERE public_template_id = %s",
                    cascade_vals,
                )
                conn.commit()
            finally:
                cur.close()
                conn.close()

        return jsonify(updated), 200

    # DELETE
    # Remove DOCX file from disk
    tmpl_dir = os.path.join(UPLOAD_DIR, str(template_id))
    if os.path.exists(tmpl_dir):
        import shutil
        shutil.rmtree(tmpl_dir)

    deleted = templates.delete(template_id)
    if not deleted:
        return jsonify({"error": "Template not found"}), 404
    return "", 204


# ============ Signature Variables API ============

@app.route("/api/templates/<int:template_id>/variables", methods=["GET", "POST"])
def api_template_variables(template_id):
    """List variables (GET) or create a variable (POST) for a template."""
    user, err = _require_auth()
    if err:
        return err

    # Verify template exists
    tmpl = templates.get_by_id(template_id)
    if tmpl is None:
        return jsonify({"error": "Template not found"}), 404

    if request.method == "GET":
        result = signature_variables.get_by_template_id(template_id)
        return jsonify(result), 200

    # POST: create variable
    data = request.get_json(silent=True)
    if data is None:
        return jsonify({"error": "Invalid JSON body"}), 400

    name = (data.get("name") or "").strip()
    if not name:
        return jsonify({"error": "Variable name is required"}), 400

    value = str(data.get("value", ""))[:1024]
    if len(str(data.get("value", ""))) > 1024:
        return jsonify({"error": "Variable value must not exceed 1024 characters"}), 400

    try:
        var = signature_variables.create(
            template_id=template_id,
            name=name,
            value=value,
            var_type=data.get("var_type", "signature"),
            page=int(data.get("page", 1)),
            x=float(data.get("x", 0)),
            y=float(data.get("y", 0)),
            width=float(data.get("width", 120)),
            height=float(data.get("height", 40)),
            font_size=int(data.get("font_size", 12)),
            font_color=str(data.get("font_color", "#000000")),
            required=bool(data.get("required", True)),
            public_variables_id=data.get("public_variables_id"),
        )
        return jsonify(var), 201
    except Exception as e:
        if "duplicate key" in str(e).lower() or "unique" in str(e).lower():
            return jsonify({"error": f"Variable name '{name}' already exists in this template"}), 409
        return jsonify({"error": str(e)}), 400


@app.route("/api/templates/<int:template_id>/variables/<int:variable_id>", methods=["PUT", "DELETE"])
def api_template_variable(template_id, variable_id):
    """Update or delete a specific variable."""
    user, err = _require_auth()
    if err:
        return err

    var = signature_variables.get_by_id(variable_id)
    if var is None or var["template_id"] != template_id:
        return jsonify({"error": "Variable not found"}), 404

    if request.method == "PUT":
        data = request.get_json(silent=True)
        if data is None:
            return jsonify({"error": "Invalid JSON body"}), 400

        updates = {}
        for field in ["name", "value", "var_type", "page", "x", "y", "width", "height",
                       "font_size", "font_color", "required", "public_variables_id"]:
            if field in data and data[field] is not None:
                val = data[field]
                if field in ("page", "font_size"):
                    val = int(val)
                elif field in ("x", "y", "width", "height"):
                    val = float(val)
                updates[field] = val

        try:
            updated = signature_variables.update(variable_id, **updates)
            if updated is None:
                return jsonify({"error": "Variable not found"}), 404
            return jsonify(updated), 200
        except Exception as e:
            if "duplicate key" in str(e).lower() or "unique" in str(e).lower():
                return jsonify({"error": "Variable name already exists in this template"}), 409
            return jsonify({"error": str(e)}), 400

    # DELETE
    signature_variables.delete(variable_id)
    return "", 204


@app.route("/api/templates/<int:template_id>/variables/import", methods=["POST"])
def api_variables_import(template_id):
    """Bulk import variables from an Excel file."""
    user, err = _require_auth()
    if err:
        return err

    tmpl = templates.get_by_id(template_id)
    if tmpl is None:
        return jsonify({"error": "Template not found"}), 404

    if "file" not in request.files:
        return jsonify({"error": "Excel file is required"}), 400

    xlsx_file = request.files["file"]
    if xlsx_file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    if not xlsx_file.filename.lower().endswith(".xlsx"):
        return jsonify({"error": "Only .xlsx files are allowed"}), 400

    try:
        import openpyxl
        wb = openpyxl.load_workbook(xlsx_file)
        ws = wb.active

        headers = [cell.value for cell in ws[1]]
        expected_headers = ["name", "value", "var_type", "page", "x", "y", "width",
                            "height", "font_size", "font_color", "required"]

        if headers[:len(expected_headers)] != expected_headers:
            return jsonify({"error": "Invalid Excel format. Expected columns: " + ", ".join(expected_headers)}), 400

        variables = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] is None:
                continue
            variables.append({
                "name": str(row[0]).strip() if row[0] is not None else "",
                "value": str(row[1]) if row[1] is not None else "",
                "var_type": str(row[2]) if row[2] is not None else "signature",
                "page": row[3] if row[3] is not None else 1,
                "x": row[4] if row[4] is not None else 0,
                "y": row[5] if row[5] is not None else 0,
                "width": row[6] if row[6] is not None else 120,
                "height": row[7] if row[7] is not None else 40,
                "font_size": row[8] if row[8] is not None else 12,
                "font_color": str(row[9]) if row[9] is not None else "#000000",
                "required": row[10] if row[10] is not None else True,
            })

        imported, errors = signature_variables.bulk_create(template_id, variables)
        return jsonify({"imported": imported, "errors": errors}), 200

    except Exception as e:
        return jsonify({"error": f"Failed to process Excel file: {str(e)}"}), 400


@app.route("/api/templates/variables-template", methods=["GET"])
def api_variables_template():
    """Download an Excel template for variable import."""
    user, err = _require_auth()
    if err:
        return err

    try:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Variables"

        headers = ["name", "value", "var_type", "page", "x", "y", "width",
                    "height", "font_size", "font_color", "required"]
        ws.append(headers)

        # Example row
        ws.append(["签字人签名", "", "signature", 1, 100, 200, 120, 40, 12, "#000000", True])

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name="variables_import_template.xlsx",
        )
    except Exception as e:
        return jsonify({"error": f"Failed to generate template: {str(e)}"}), 500


# ============ Public Variables (template-independent) ============

@app.route("/api/public-variables", methods=["GET"])
def api_public_variables_list():
    """List all public variables."""
    user, err = _require_auth()
    if err:
        return err
    return jsonify(public_variables.get_all()), 200


@app.route("/api/public-variables", methods=["POST"])
def api_public_variables_create():
    """Create a new public variable."""
    user, err = _require_auth()
    if err:
        return err

    data = request.get_json(silent=True)
    if data is None:
        return jsonify({"error": "Invalid JSON body"}), 400

    name = data.get("name", "").strip()
    if not name:
        return jsonify({"error": "Variable name is required"}), 400

    kwargs = {
        "name": name,
        "value": data.get("value", ""),
        "var_type": data.get("var_type", "signature"),
        "page": int(data.get("page", 1)),
        "x": float(data.get("x", 0)),
        "y": float(data.get("y", 0)),
        "width": float(data.get("width", 120)),
        "height": float(data.get("height", 40)),
        "font_size": int(data.get("font_size", 12)),
        "font_color": data.get("font_color", "#000000"),
        "required": data.get("required", True),
    }

    try:
        created = public_variables.create(**kwargs)
        return jsonify(created), 201
    except Exception as e:
        if "duplicate" in str(e).lower() or "unique" in str(e).lower():
            return jsonify({"error": "Variable name already exists"}), 409
        return jsonify({"error": str(e)}), 400


@app.route("/api/public-variables/<int:variable_id>", methods=["PUT"])
def api_public_variables_update(variable_id):
    """Update a public variable."""
    user, err = _require_auth()
    if err:
        return err

    data = request.get_json(silent=True)
    if data is None:
        return jsonify({"error": "Invalid JSON body"}), 400

    updates = {}
    for field in ["name", "value", "var_type", "page", "x", "y", "width", "height",
                   "font_size", "font_color", "required"]:
        if field in data and data[field] is not None:
            val = data[field]
            if field in ("page", "font_size"):
                val = int(val)
            elif field in ("x", "y", "width", "height"):
                val = float(val)
            updates[field] = val

    try:
        updated = public_variables.update(variable_id, **updates)
        if updated is None:
            return jsonify({"error": "Variable not found"}), 404
        return jsonify(updated), 200
    except Exception as e:
        if "duplicate" in str(e).lower() or "unique" in str(e).lower():
            return jsonify({"error": "Variable name already exists"}), 409
        return jsonify({"error": str(e)}), 400


@app.route("/api/public-variables/<int:variable_id>", methods=["DELETE"])
def api_public_variables_delete(variable_id):
    """Delete a public variable."""
    user, err = _require_auth()
    if err:
        return err

    deleted = public_variables.delete(variable_id)
    if not deleted:
        return jsonify({"error": "Variable not found"}), 404
    return "", 204


@app.route("/api/public-variables/import", methods=["POST"])
def api_public_variables_import():
    """Bulk import public variables from an Excel file."""
    user, err = _require_auth()
    if err:
        return err

    if "file" not in request.files:
        return jsonify({"error": "Excel file is required"}), 400

    xlsx_file = request.files["file"]
    if xlsx_file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    if not xlsx_file.filename.lower().endswith(".xlsx"):
        return jsonify({"error": "Only .xlsx files are allowed"}), 400

    try:
        import openpyxl
        wb = openpyxl.load_workbook(xlsx_file)
        ws = wb.active

        headers = [cell.value for cell in ws[1]]
        expected_headers = ["name", "value", "var_type", "page", "x", "y", "width",
                            "height", "font_size", "font_color", "required"]

        if headers[:len(expected_headers)] != expected_headers:
            return jsonify({"error": "Invalid Excel format. Expected columns: " + ", ".join(expected_headers)}), 400

        variables = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] is None:
                continue
            variables.append({
                "name": str(row[0]).strip() if row[0] is not None else "",
                "value": str(row[1]) if row[1] is not None else "",
                "var_type": str(row[2]) if row[2] is not None else "signature",
                "page": row[3] if row[3] is not None else 1,
                "x": row[4] if row[4] is not None else 0,
                "y": row[5] if row[5] is not None else 0,
                "width": row[6] if row[6] is not None else 120,
                "height": row[7] if row[7] is not None else 40,
                "font_size": row[8] if row[8] is not None else 12,
                "font_color": str(row[9]) if row[9] is not None else "#000000",
                "required": row[10] if row[10] is not None else True,
            })

        imported, errors = public_variables.bulk_create(variables)
        return jsonify({"imported": imported, "errors": errors}), 200

    except Exception as e:
        return jsonify({"error": f"Failed to process Excel file: {str(e)}"}), 400


@app.route("/api/public-variables/import-template", methods=["GET"])
def api_public_variables_template():
    """Download an Excel template for public variables import."""
    user, err = _require_auth()
    if err:
        return err

    try:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Public Variables"

        headers = ["name", "value", "var_type", "page", "x", "y", "width",
                    "height", "font_size", "font_color", "required"]
        ws.append(headers)

        # Example row
        ws.append(["甲方签字", "", "signature", 1, 100, 200, 120, 40, 12, "#000000", True])

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name="public_variables_import_template.xlsx",
        )
    except Exception as e:
        return jsonify({"error": f"Failed to generate template: {str(e)}"}), 500


# ============ Template Selection ============

@app.route("/api/projects/<int:project_id>/select-template", methods=["POST"])
def api_select_template(project_id):
    """Associate a template with a project."""
    user, err = _require_auth()
    if err:
        return err

    project = projects.get_by_id(project_id)
    if project is None:
        return jsonify({"error": "Project not found"}), 404

    if project["created_by"] != user.username:
        return jsonify({"error": "Forbidden"}), 403

    data = request.get_json(silent=True)
    if data is None:
        return jsonify({"error": "Invalid JSON body"}), 400

    template_id = data.get("template_id")
    if not template_id:
        return jsonify({"error": "template_id is required"}), 400

    # Verify template exists
    tmpl = templates.get_by_id(template_id)
    if tmpl is None:
        return jsonify({"error": "Template not found"}), 404

    updated = projects.update(project_id, template_id=template_id, status="template_selected")
    return jsonify(updated), 200


# ============ Project-Scoped Template Management ============

@app.route("/api/projects/<int:project_id>/templates", methods=["GET"])
def api_project_templates_list(project_id):
    """List templates for a specific project, optionally filtered by category."""
    user, err = _require_auth()
    if err:
        return err

    proj = projects.get_by_id(project_id)
    if proj is None:
        return jsonify({"error": "Project not found"}), 404

    category = request.args.get("category")
    templates_list = project_templates.list_all(project_id, category=category)
    return jsonify(templates_list), 200


@app.route("/api/projects/<int:project_id>/templates", methods=["POST"])
def api_project_templates_create(project_id):
    """Upload a DOCX template for a specific project."""
    user, err = _require_auth()
    if err:
        return err

    proj = projects.get_by_id(project_id)
    if proj is None:
        return jsonify({"error": "Project not found"}), 404

    name = request.form.get("name", "").strip()
    category = request.form.get("category", "").strip()
    description = request.form.get("description", "").strip()

    if not name:
        return jsonify({"error": "Template name is required"}), 400
    if not category:
        return jsonify({"error": "Category is required"}), 400

    if "file" not in request.files:
        return jsonify({"error": "DOCX file is required"}), 400

    docx_file = request.files["file"]
    if docx_file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    if not docx_file.filename.lower().endswith(".docx"):
        return jsonify({"error": "Only DOCX files are allowed"}), 400

    # Create template record first to get id
    tmpl = project_templates.create(project_id, name, category, description, "")
    template_id = tmpl["id"]

    # Save DOCX file
    upload_dir = os.path.join("uploads", "project_templates", str(template_id))
    os.makedirs(upload_dir, exist_ok=True)
    file_path = os.path.join(upload_dir, f"{template_id}.docx")
    docx_file.save(file_path)

    # Update file_path in record
    tmpl = project_templates.update(template_id, file_path=file_path)
    return jsonify(tmpl), 201


@app.route("/api/projects/<int:project_id>/templates/<int:template_id>/preview", methods=["GET"])
def api_project_template_preview(project_id, template_id):
    """Generate a preview PDF: replace ${varName} in DOCX with variable values, convert to PDF."""
    user, err = _require_auth()
    if err:
        return err

    tmpl = project_templates.get_by_id(template_id)
    if tmpl is None:
        return jsonify({"error": "Template not found"}), 404

    docx_path = tmpl.get("file_path")
    if not docx_path or not os.path.exists(docx_path):
        return jsonify({"error": "Template DOCX file not found on disk"}), 404

    # Gather project template variables (skip empty/null values to preserve placeholders)
    proj_vars = project_variables.get_by_template_id(template_id)
    variables: dict = {}
    for v in proj_vars:
        val = v.get("value")
        if val:  # only include non-empty values; empty/null → keep placeholder
            variables[v["name"]] = str(val)

    # Gather public template variables (if bound to a public template)
    public_tmpl_id = tmpl.get("public_template_id")
    if public_tmpl_id:
        pub_vars = signature_variables.get_by_template_id(int(public_tmpl_id))
        # Merge: project variables override public variables by name
        for v in pub_vars:
            name = v["name"]
            if name not in variables:
                val = v.get("value")
                if val:  # only include non-empty values; empty/null → keep placeholder
                    variables[name] = str(val)

    # ---------- DOCX variable replacement ----------
    import re
    from docx import Document

    doc = Document(docx_path)

    if variables:
        def _replace_in_paragraph(paragraph, repl_map):
            """Replace ${varName} tokens in a paragraph, handling cross-run splitting."""
            # Pass 1: per-run replacement (handles most cases)
            for run in paragraph.runs:
                text = run.text
                for name, value in repl_map.items():
                    text = text.replace(name, value)
                run.text = text

            # Pass 2: check if any pattern still spans multiple runs
            remaining = re.search(r'\$\{[^}]+\}', paragraph.text)
            if remaining:
                full_text = paragraph.text
                for name, value in repl_map.items():
                    full_text = full_text.replace(name, value)
                if paragraph.runs:
                    ref_run = paragraph.runs[0]
                    paragraph.clear()
                    new_run = paragraph.add_run(full_text)
                    if ref_run.font.name:
                        new_run.font.name = ref_run.font.name
                    if ref_run.font.size:
                        new_run.font.size = ref_run.font.size
                    new_run.bold = ref_run.bold
                    new_run.italic = ref_run.italic
                else:
                    paragraph.text = full_text

        for paragraph in doc.paragraphs:
            _replace_in_paragraph(paragraph, variables)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        _replace_in_paragraph(paragraph, variables)

    # ---------- Save temp DOCX and convert to PDF ----------
    tmpl_dir = os.path.dirname(docx_path)
    temp_docx = os.path.join(tmpl_dir, "_preview.docx")
    preview_pdf = os.path.join(tmpl_dir, "preview.pdf")

    doc.save(temp_docx)

    # Delete old preview PDF if exists
    if os.path.exists(preview_pdf):
        os.remove(preview_pdf)

    try:
        from docx2pdf import convert
        convert(temp_docx, preview_pdf)
    except Exception as e:
        # Clean up temp file and return error
        if os.path.exists(temp_docx):
            os.remove(temp_docx)
        return jsonify({"error": f"PDF conversion failed: {str(e)}"}), 500
    finally:
        # Always clean up temp DOCX
        if os.path.exists(temp_docx):
            os.remove(temp_docx)

    return send_file(preview_pdf, mimetype="application/pdf")


@app.route("/api/projects/<int:project_id>/templates/<int:template_id>", methods=["GET"])
def api_project_templates_get(project_id, template_id):
    """Get a single project template by its ID."""
    user, err = _require_auth()
    if err:
        return err

    tmpl = project_templates.get_by_id(template_id)
    if tmpl is None:
        return jsonify({"error": "Template not found"}), 404

    return jsonify(tmpl), 200


@app.route("/api/projects/<int:project_id>/templates/<int:template_id>", methods=["PUT"])
def api_project_templates_update(project_id, template_id):
    """Update a project template's metadata (name, category, description)."""
    user, err = _require_auth()
    if err:
        return err

    proj = projects.get_by_id(project_id)
    if proj is None:
        return jsonify({"error": "Project not found"}), 404

    tmpl = project_templates.get_by_id(template_id)
    if tmpl is None:
        return jsonify({"error": "Template not found"}), 404

    data = request.get_json(silent=True)
    if data is None:
        return jsonify({"error": "Invalid JSON body"}), 400

    name = data.get("name")
    category = data.get("category")
    description = data.get("description")

    updated = project_templates.update(template_id, name=name, category=category, description=description)
    return jsonify(updated), 200


@app.route("/api/projects/<int:project_id>/templates/<int:template_id>", methods=["DELETE"])
def api_project_templates_delete(project_id, template_id):
    """Delete a project template and its associated DOCX file."""
    user, err = _require_auth()
    if err:
        return err

    proj = projects.get_by_id(project_id)
    if proj is None:
        return jsonify({"error": "Project not found"}), 404

    tmpl = project_templates.get_by_id(template_id)
    if tmpl is None:
        return jsonify({"error": "Template not found"}), 404

    # Remove DOCX file if it exists
    if tmpl.get("file_path"):
        file_path = tmpl["file_path"]
        if os.path.exists(file_path):
            os.remove(file_path)
        # Remove parent directory if empty
        parent_dir = os.path.dirname(file_path)
        try:
            if os.path.exists(parent_dir) and not os.listdir(parent_dir):
                os.rmdir(parent_dir)
        except OSError:
            pass

    project_templates.delete(template_id)
    return "", 204


@app.route("/api/projects/<int:project_id>/templates/bind", methods=["POST"])
def api_project_templates_bind(project_id):
    """Bind (deep-copy) one or more public templates into a project."""
    user, err = _require_auth()
    if err:
        return err

    proj = projects.get_by_id(project_id)
    if proj is None:
        return jsonify({"error": "Project not found"}), 404

    body = request.get_json(silent=True)
    if body is None:
        return jsonify({"error": "Invalid JSON body"}), 400

    template_ids = body.get("template_ids")
    if not template_ids or not isinstance(template_ids, list):
        return jsonify({"error": "template_ids list is required"}), 400

    import shutil

    conn = get_connection()
    cur = conn.cursor()
    bound = []
    skipped = []

    try:
        for tid in template_ids:
            pub_tmpl = templates.get_by_id(int(tid))
            if pub_tmpl is None:
                skipped.append({"template_id": tid, "reason": "模板不存在"})
                continue

            # Check duplicate by (project_id, name)
            cur.execute(
                "SELECT id FROM project_templates WHERE project_id = %s AND name = %s",
                (project_id, pub_tmpl["name"]),
            )
            if cur.fetchone():
                skipped.append({
                    "template_id": tid,
                    "name": pub_tmpl["name"],
                    "reason": "已存在同名模板",
                })
                continue

            # 1) Create project template record
            cur.execute(
                """INSERT INTO project_templates
                   (project_id, name, category, file_path, description, public_template_id)
                   VALUES (%s, %s, %s, %s, %s, %s)
                   RETURNING id, project_id, name, category, file_path, description, public_template_id""",
                (
                    project_id,
                    pub_tmpl["name"],
                    pub_tmpl.get("category", ""),
                    "",
                    pub_tmpl.get("description", ""),
                    pub_tmpl["id"],
                ),
            )
            row = cur.fetchone()
            cols = [desc[0] for desc in cur.description]
            new_tmpl = dict(zip(cols, row))
            new_id = new_tmpl["id"]

            # 2) Copy DOCX file
            src_path = pub_tmpl.get("file_path")
            if src_path and os.path.exists(src_path):
                dest_dir = os.path.join("uploads", "project_templates", str(new_id))
                os.makedirs(dest_dir, exist_ok=True)
                dest_path = os.path.join(dest_dir, f"{new_id}.docx")
                shutil.copy2(src_path, dest_path)
                cur.execute(
                    "UPDATE project_templates SET file_path = %s WHERE id = %s",
                    (dest_path, new_id),
                )
                new_tmpl["file_path"] = dest_path

            # 3) Copy variables
            pub_vars = signature_variables.get_by_template_id(int(tid))
            for var in pub_vars:
                cur.execute(
                    """INSERT INTO project_template_variables
                       (template_id, name, value, var_type, page, x, y,
                        width, height, font_size, font_color, required, public_variables_id)
                       VALUES (%s, %s, %s, %s, %s, %s, %s,
                               %s, %s, %s, %s, %s, %s)""",
                    (
                        new_id,
                        var["name"],
                        var.get("value", ""),
                        var.get("var_type", "signature"),
                        var.get("page", 1),
                        var.get("x", 0),
                        var.get("y", 0),
                        var.get("width", 120),
                        var.get("height", 40),
                        var.get("font_size", 12),
                        var.get("font_color", "#000000"),
                        var.get("required", True),
                        var.get("public_variables_id"),
                    ),
                )

            bound.append({
                "template_id": tid,
                "new_id": new_id,
                "name": new_tmpl["name"],
            })

        conn.commit()
        return jsonify({"bound": bound, "skipped": skipped}), 201

    except Exception as exc:
        conn.rollback()
        print(f"[bind] ERROR: {exc}", file=sys.stderr)
        return jsonify({"error": f"绑定失败: {exc}"}), 500
    finally:
        cur.close()
        conn.close()


# ============ Project-Scoped Template Variables ============

@app.route("/api/projects/<int:project_id>/templates/<int:template_id>/variables", methods=["GET"])
def api_project_variables_list(project_id, template_id):
    """List all variables for a project template."""
    user, err = _require_auth()
    if err:
        return err

    proj = projects.get_by_id(project_id)
    if proj is None:
        return jsonify({"error": "Project not found"}), 404

    tmpl = project_templates.get_by_id(template_id)
    if tmpl is None:
        return jsonify({"error": "Template not found"}), 404

    variables = project_variables.get_by_template_id(template_id)
    return jsonify(variables), 200


@app.route("/api/projects/<int:project_id>/templates/<int:template_id>/variables", methods=["POST"])
def api_project_variables_create(project_id, template_id):
    """Create a new variable for a project template."""
    user, err = _require_auth()
    if err:
        return err

    proj = projects.get_by_id(project_id)
    if proj is None:
        return jsonify({"error": "Project not found"}), 404

    tmpl = project_templates.get_by_id(template_id)
    if tmpl is None:
        return jsonify({"error": "Template not found"}), 404

    data = request.get_json(silent=True)
    if data is None:
        return jsonify({"error": "Invalid JSON body"}), 400

    name = data.get("name", "").strip()
    if not name:
        return jsonify({"error": "Variable name is required"}), 400

    value = str(data.get("value", ""))[:1024]
    if len(str(data.get("value", ""))) > 1024:
        return jsonify({"error": "Variable value must not exceed 1024 characters"}), 400

    try:
        var = project_variables.create(
            template_id=template_id,
            name=name,
            value=value,
            var_type=data.get("var_type", "signature"),
            page=data.get("page", 1),
            x=data.get("x", 0),
            y=data.get("y", 0),
            width=data.get("width", 120),
            height=data.get("height", 40),
            font_size=data.get("font_size", 12),
            font_color=data.get("font_color", "#000000"),
            required=data.get("required", True),
            public_variables_id=data.get("public_variables_id"),
        )
        return jsonify(var), 201
    except Exception as e:
        if "duplicate" in str(e).lower() or "unique" in str(e).lower():
            return jsonify({"error": f"Variable '{name}' already exists in this template"}), 409
        return jsonify({"error": str(e)}), 500


@app.route("/api/projects/<int:project_id>/templates/<int:template_id>/variables/<int:variable_id>", methods=["PUT"])
def api_project_variables_update(project_id, template_id, variable_id):
    """Update a project template variable."""
    user, err = _require_auth()
    if err:
        return err

    proj = projects.get_by_id(project_id)
    if proj is None:
        return jsonify({"error": "Project not found"}), 404

    tmpl = project_templates.get_by_id(template_id)
    if tmpl is None:
        return jsonify({"error": "Template not found"}), 404

    var = project_variables.get_by_id(variable_id)
    if var is None:
        return jsonify({"error": "Variable not found"}), 404

    data = request.get_json(silent=True)
    if data is None:
        return jsonify({"error": "Invalid JSON body"}), 400

    update_kwargs = {}
    for field in ["name", "value", "var_type", "page", "x", "y", "width", "height",
                  "font_size", "font_color", "required", "public_variables_id"]:
        if field in data:
            update_kwargs[field] = data[field]

    updated = project_variables.update(variable_id, **update_kwargs)
    return jsonify(updated), 200


@app.route("/api/projects/<int:project_id>/templates/<int:template_id>/variables/<int:variable_id>", methods=["DELETE"])
def api_project_variables_delete(project_id, template_id, variable_id):
    """Delete a project template variable."""
    user, err = _require_auth()
    if err:
        return err

    proj = projects.get_by_id(project_id)
    if proj is None:
        return jsonify({"error": "Project not found"}), 404

    tmpl = project_templates.get_by_id(template_id)
    if tmpl is None:
        return jsonify({"error": "Template not found"}), 404

    var = project_variables.get_by_id(variable_id)
    if var is None:
        return jsonify({"error": "Variable not found"}), 404

    project_variables.delete(variable_id)
    return "", 204


@app.route("/api/projects/<int:project_id>/templates/<int:template_id>/variables/import", methods=["POST"])
def api_project_variables_import(project_id, template_id):
    """Bulk import variables from an Excel file for a project template."""
    user, err = _require_auth()
    if err:
        return err

    proj = projects.get_by_id(project_id)
    if proj is None:
        return jsonify({"error": "Project not found"}), 404

    tmpl = project_templates.get_by_id(template_id)
    if tmpl is None:
        return jsonify({"error": "Template not found"}), 404

    if "file" not in request.files:
        return jsonify({"error": "Excel file is required"}), 400

    xlsx_file = request.files["file"]
    if xlsx_file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    if not xlsx_file.filename.lower().endswith(".xlsx"):
        return jsonify({"error": "Only .xlsx files are allowed"}), 400

    try:
        import openpyxl
        wb = openpyxl.load_workbook(xlsx_file)
        ws = wb.active

        headers = [cell.value for cell in ws[1]]
        expected_headers = ["name", "value", "var_type", "page", "x", "y", "width",
                            "height", "font_size", "font_color", "required"]

        if headers[:len(expected_headers)] != expected_headers:
            return jsonify({"error": "Invalid Excel format. Expected columns: " + ", ".join(expected_headers)}), 400

        variables = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row[0] is None:
                continue
            variables.append({
                "name": str(row[0]).strip() if row[0] is not None else "",
                "value": str(row[1]) if row[1] is not None else "",
                "var_type": str(row[2]) if row[2] is not None else "signature",
                "page": row[3] if row[3] is not None else 1,
                "x": row[4] if row[4] is not None else 0,
                "y": row[5] if row[5] is not None else 0,
                "width": row[6] if row[6] is not None else 120,
                "height": row[7] if row[7] is not None else 40,
                "font_size": row[8] if row[8] is not None else 12,
                "font_color": str(row[9]) if row[9] is not None else "#000000",
                "required": row[10] if row[10] is not None else True,
            })

        imported, errors = project_variables.bulk_create(template_id, variables)
        return jsonify({"imported": imported, "errors": errors}), 200

    except Exception as e:
        return jsonify({"error": f"Failed to process Excel file: {str(e)}"}), 400


@app.route("/api/projects/<int:project_id>/templates/variables-template", methods=["GET"])
def api_project_variables_template(project_id):
    """Download an Excel template for project variable import."""
    user, err = _require_auth()
    if err:
        return err

    proj = projects.get_by_id(project_id)
    if proj is None:
        return jsonify({"error": "Project not found"}), 404

    try:
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Variables"

        headers = ["name", "value", "var_type", "page", "x", "y", "width",
                    "height", "font_size", "font_color", "required"]
        ws.append(headers)

        # Example row
        ws.append(["签字人签名", "", "signature", 1, 100, 200, 120, 40, 12, "#000000", True])

        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        return send_file(
            output,
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            as_attachment=True,
            download_name="project_variables_import_template.xlsx",
        )
    except Exception as e:
        return jsonify({"error": f"Failed to generate template: {str(e)}"}), 500


if __name__ == "__main__":
    app.run(debug=True, host="127.0.0.1", port=8090)
